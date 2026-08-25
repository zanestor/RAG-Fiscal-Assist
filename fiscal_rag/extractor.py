from __future__ import annotations

import hashlib
import os
import re
import shutil
from pathlib import Path
from typing import Any


MIN_PAGE_CHARACTERS = 40
MIN_DOCUMENT_CHARACTERS = 200
HEADER_END_MARKER = "Do not infer legal force, effective date, amendment, or repeal status from the catalog date alone."


def ocr_configuration() -> tuple[str, Path]:
    languages = os.getenv("FISCAL_RAG_OCR_LANGUAGES", "fra+eng").strip() or "fra+eng"
    configured = os.getenv("FISCAL_RAG_TESSDATA_DIR") or os.getenv("TESSDATA_PREFIX")
    if configured:
        tessdata_dir = Path(os.path.expandvars(configured)).expanduser()
    else:
        executable = shutil.which("tesseract")
        tessdata_dir = Path(executable).resolve().parent / "tessdata" if executable else Path("tessdata")
    return languages, tessdata_dir


def validate_ocr_configuration() -> tuple[str, Path]:
    languages, tessdata_dir = ocr_configuration()
    requested = [language for language in languages.split("+") if language]
    missing = [language for language in requested if not (tessdata_dir / f"{language}.traineddata").is_file()]
    if missing:
        names = ", ".join(f"{language}.traineddata" for language in missing)
        raise RuntimeError(
            f"OCR language data missing from {tessdata_dir}: {names}. "
            "Run: powershell -ExecutionPolicy Bypass -File .\\setup_ocr_languages.ps1"
        )
    return languages, tessdata_dir


def sha256_file(path: Path, block_size: int = 1024 * 1024) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        while block := handle.read(block_size):
            digest.update(block)
    return digest.hexdigest()


def safe_filename(value: str, limit: int = 72) -> str:
    normalized = re.sub(r"[^A-Za-z0-9._-]+", "_", value).strip("._")
    return (normalized or "document")[:limit]


def header_lines(document: dict[str, Any]) -> list[str]:
    lines = [
        "# Fiscal reference source",
        "",
        f"Title: {document['title']}",
        f"Source institution: {document['source_label']}",
        f"Source ID: {document['source']}",
        f"Document ID: {document['id']}",
        f"Original filename: {document['filename']}",
        f"Repository path: {document['relative_path']}",
    ]
    if document.get("category"):
        lines.append(f"Category: {document['category']}")
    if document.get("published_date"):
        lines.append(f"Catalog date: {document['published_date']}")
    if document.get("source_url"):
        lines.append(f"Original URL: {document['source_url']}")
    lines.extend(
        [
            "",
            HEADER_END_MARKER,
        ]
    )
    return lines


def refresh_extraction_header(document: dict[str, Any], extracted_path: Path) -> bool:
    """Update catalog metadata without touching the extracted/OCR evidence body.

    Returns ``False`` for an unfamiliar legacy file so the caller can fall back to
    a normal extraction. Every extraction created by this pipeline starts with the
    synthetic header and fixed trailer validated here.
    """
    try:
        content = extracted_path.read_text(encoding="utf-8", errors="replace")
    except OSError:
        return False
    if not content.startswith("# Fiscal reference source"):
        return False
    marker_start = content.find(HEADER_END_MARKER)
    if marker_start == -1:
        return False
    body_start = marker_start + len(HEADER_END_MARKER)
    updated = "\n".join(header_lines(document)) + content[body_start:]
    if updated != content:
        extracted_path.write_text(updated, encoding="utf-8")
    return True


def output_path_for(document: dict[str, Any], output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    output_name = f"{document['id']}__{safe_filename(Path(document['filename']).stem)}.md"
    return output_dir / output_name


def write_text_extraction(
    document: dict[str, Any],
    output_dir: Path,
    sections: list[str],
    locator_note: str,
) -> dict[str, Any]:
    output_path = output_path_for(document, output_dir)
    content = "\n".join(header_lines(document) + ["", locator_note]) + "".join(sections) + "\n"
    output_path.write_text(content, encoding="utf-8")
    character_count = sum(len(section) for section in sections)
    return {
        "status": "ready" if character_count >= MIN_DOCUMENT_CHARACTERS else "insufficient_text",
        "extracted_path": str(output_path),
        "extracted_filename": output_path.name,
        "page_count": 0,
        "character_count": character_count,
        "ocr_pages": [],
        "empty_pages": [],
        "warnings": [],
    }


def extract_pdf(document: dict[str, Any], output_dir: Path, use_ocr: bool = False, force_ocr: bool = False) -> dict[str, Any]:
    try:
        import fitz  # PyMuPDF
    except ImportError as error:
        raise RuntimeError("PyMuPDF is required. Run: pip install -r requirements.txt") from error

    pdf_path = Path(document["absolute_path"])
    output_path = output_path_for(document, output_dir)

    pages: list[str] = []
    text_character_count = 0
    ocr_pages: list[int] = []
    empty_pages: list[int] = []
    extraction_errors: list[str] = []

    try:
        pdf = fitz.open(pdf_path)
    except Exception as error:
        return {
            "status": "error",
            "error": f"Unable to open PDF: {error}",
            "page_count": 0,
            "character_count": 0,
            "ocr_pages": [],
            "empty_pages": [],
        }

    try:
        ocr_languages = ""
        tessdata_dir: Path | None = None
        if use_ocr:
            ocr_languages, tessdata_dir = validate_ocr_configuration()
        for page_index, page in enumerate(pdf):
            page_number = page_index + 1
            text = page.get_text("text", sort=True).strip()
            # force_ocr exists for a different failure mode than the normal
            # "page has too little text" check below: a PDF whose embedded
            # text layer is long enough to pass that check but is garbled
            # (duplicated words, mis-encoded accents - seen in a pre-OCR'd
            # source file added to this corpus) never gets re-OCR'd otherwise,
            # since length alone looks fine. force_ocr skips that length gate
            # and, since the existing text is known bad, trusts the fresh OCR
            # result outright rather than only keeping it if it's longer -
            # duplicated garbage text is often LONGER than the correct text,
            # so a length comparison would wrongly keep the garbage.
            if (force_ocr or len(text) < MIN_PAGE_CHARACTERS) and use_ocr:
                pixel_width = round(page.rect.width * 150 / 72)
                pixel_height = round(page.rect.height * 150 / 72)
                if pixel_width < 3 or pixel_height < 3:
                    extraction_errors.append(
                        f"OCR skipped page {page_number}: rendered page is too small ({pixel_width}x{pixel_height})."
                    )
                else:
                    try:
                        text_page = page.get_textpage_ocr(
                            language=ocr_languages,
                            dpi=150,
                            full=True,
                            tessdata=str(tessdata_dir),
                        )
                        ocr_text = page.get_text("text", textpage=text_page, sort=True).strip()
                        # Trust a forced OCR result outright (even if shorter than
                        # duplicated-garbage embedded text) UNLESS it's so short it
                        # looks like OCR itself failed (missing tessdata language, a
                        # blank scan) - in that case keep the existing text rather
                        # than replacing something readable with near-nothing.
                        if (force_ocr and (len(ocr_text) >= MIN_PAGE_CHARACTERS or len(ocr_text) > len(text))) or (
                            not force_ocr and len(ocr_text) > len(text)
                        ):
                            text = ocr_text
                            ocr_pages.append(page_number)
                    except Exception as error:
                        extraction_errors.append(f"OCR page {page_number}: {error}")
            if len(text) < MIN_PAGE_CHARACTERS:
                empty_pages.append(page_number)
            text_character_count += len(text)
            pages.append(f"\n\n<!-- PAGE {page_number} -->\n## Page {page_number}\n\n{text or '[No extractable text on this page]'}")
    finally:
        pdf.close()

    character_count = text_character_count
    metadata = header_lines(document)
    metadata.append("Important: page headings below correspond to the original PDF page numbers.")
    output_path.write_text("\n".join(metadata) + "".join(pages) + "\n", encoding="utf-8")
    status = "ready" if character_count >= MIN_DOCUMENT_CHARACTERS else "needs_ocr"
    return {
        "status": status,
        "extracted_path": str(output_path),
        "extracted_filename": output_path.name,
        "page_count": len(pages),
        "character_count": character_count,
        "ocr_pages": ocr_pages,
        "empty_pages": empty_pages,
        "warnings": extraction_errors,
    }


def read_text_file(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-16", "cp1252"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def extract_plain_text(document: dict[str, Any], output_dir: Path) -> dict[str, Any]:
    text = read_text_file(Path(document["absolute_path"])).strip()
    sections = [f"\n\n<!-- SECTION 1 -->\n## Section 1\n\n{text or '[No extractable text]'}"]
    return write_text_extraction(
        document,
        output_dir,
        sections,
        "Section headings below identify stable text sections in the original file.",
    )


def extract_docx(document: dict[str, Any], output_dir: Path) -> dict[str, Any]:
    try:
        from docx import Document as WordDocument
    except ImportError as error:
        raise RuntimeError("python-docx is required. Run: pip install -r requirements.txt") from error
    word_document = WordDocument(document["absolute_path"])
    blocks: list[str] = [paragraph.text.strip() for paragraph in word_document.paragraphs if paragraph.text.strip()]
    for table_index, table in enumerate(word_document.tables, start=1):
        blocks.append(f"Table {table_index}")
        for row_index, row in enumerate(table.rows, start=1):
            values = [cell.text.replace("\n", " ").strip() for cell in row.cells]
            blocks.append(f"Row {row_index}: " + " | ".join(values))
    body = "\n\n".join(blocks) or "[No extractable text]"
    sections = [f"\n\n<!-- SECTION 1 -->\n## Section 1\n\n{body}"]
    return write_text_extraction(document, output_dir, sections, "Section and table-row markers identify locations in the Word document.")


def extract_xlsx(document: dict[str, Any], output_dir: Path) -> dict[str, Any]:
    try:
        import openpyxl
    except ImportError as error:
        raise RuntimeError("openpyxl is required. Run: pip install -r requirements.txt") from error
    workbook = openpyxl.load_workbook(document["absolute_path"], read_only=True, data_only=True)
    sections: list[str] = []
    try:
        for worksheet in workbook.worksheets:
            lines: list[str] = []
            for row_number, row in enumerate(worksheet.iter_rows(values_only=True), start=1):
                values = ["" if value is None else str(value).replace("\n", " ").strip() for value in row]
                if any(values):
                    lines.append(f"Row {row_number}: " + " | ".join(values))
            sections.append(f"\n\n<!-- SHEET {worksheet.title} -->\n## Sheet: {worksheet.title}\n\n" + ("\n".join(lines) or "[No populated rows]"))
    finally:
        workbook.close()
    return write_text_extraction(document, output_dir, sections, "Sheet names and row numbers below identify locations in the original workbook.")


def extract_xls(document: dict[str, Any], output_dir: Path) -> dict[str, Any]:
    try:
        import xlrd
    except ImportError as error:
        raise RuntimeError("xlrd is required for .xls files. Run: pip install -r requirements.txt") from error
    workbook = xlrd.open_workbook(document["absolute_path"], on_demand=True)
    sections: list[str] = []
    try:
        for worksheet in workbook.sheets():
            lines = []
            for row_index in range(worksheet.nrows):
                values = [str(value).replace("\n", " ").strip() for value in worksheet.row_values(row_index)]
                if any(values):
                    lines.append(f"Row {row_index + 1}: " + " | ".join(values))
            sections.append(f"\n\n<!-- SHEET {worksheet.name} -->\n## Sheet: {worksheet.name}\n\n" + ("\n".join(lines) or "[No populated rows]"))
    finally:
        workbook.release_resources()
    return write_text_extraction(document, output_dir, sections, "Sheet names and row numbers below identify locations in the original workbook.")


def extract_document(document: dict[str, Any], output_dir: Path, use_ocr: bool = False, force_ocr: bool = False) -> dict[str, Any]:
    extension = Path(document["absolute_path"]).suffix.lower()
    if extension == ".pdf":
        return extract_pdf(document, output_dir, use_ocr=use_ocr, force_ocr=force_ocr)
    if extension == ".docx":
        return extract_docx(document, output_dir)
    if extension == ".xlsx":
        return extract_xlsx(document, output_dir)
    if extension == ".xls":
        return extract_xls(document, output_dir)
    if extension in {".txt", ".md", ".csv"}:
        return extract_plain_text(document, output_dir)
    raise RuntimeError(f"Unsupported document type: {extension}")
